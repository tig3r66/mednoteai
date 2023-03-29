import openai
import whisper
import pickle
import os
import random

from moviepy.editor import *
from pydub import AudioSegment
import math
import os
import glob
import docx
import tqdm

import PyPDF2


from dotenv import load_dotenv

load_dotenv()
openai.api_key = os.getenv('API_KEY')

model = whisper.load_model("base")


def MP4ToMP3(mp4, mp3):
    FILETOCONVERT = AudioFileClip(mp4)
    FILETOCONVERT.write_audiofile(mp3)
    FILETOCONVERT.close()


def split_audio(fname, output_dir='memory/chunks', chunk_duration=3):
    # chunk duration in minutes
    chunk_duration = chunk_duration * 60 * 1000
    sound = AudioSegment.from_file(f'memory/{fname}')
    num_chunks = math.ceil(len(sound) / chunk_duration)
    os.makedirs(output_dir, exist_ok=True)
    for i in range(num_chunks):
        start = i * chunk_duration
        end = min((i + 1) * chunk_duration, len(sound))
        chunk = sound[start:end]
        chunk.export(f"memory/chunks/{fname}_chunk{i}.mp3", format="mp3")


def split_audio_audio(fname, output_dir='memory/chunks', chunk_duration=3):
    # chunk duration in minutes
    chunk_duration = chunk_duration * 60 * 1000
    sound = AudioSegment.from_file(f'static/upload/{fname}')
    num_chunks = math.ceil(len(sound) / chunk_duration)
    os.makedirs(output_dir, exist_ok=True)
    for i in tqdm.tqdm(range(num_chunks)):
        start = i * chunk_duration
        end = min((i + 1) * chunk_duration, len(sound))
        chunk = sound[start:end]
        chunk.export(f"memory/chunks/{fname}_chunk{i}.mp3", format="mp3")


def pdf_to_content(fname):
    page_content = ''
    with open(fname, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        num_pages = len(pdf_reader.pages)
        for page_num in range(1, num_pages):
            try:
                page_content += pdf_reader.pages[page_num].extract_text()
            except:
                continue

    max_length = 3000
    content = []
    start = 0
    while start < len(page_content):
        end = start + max_length
        chunk = page_content[start:end]
        content.append(chunk)
        start = end

    return content


def make_note(fname, is_transcribed):
    if is_transcribed is False:
        if fname.endswith('.mp4'):
            fname = fname.strip('.mp4')
            MP4ToMP3(f'static/upload/{fname}.mp4', f'memory/{fname}.mp3')
            split_audio(f'{fname}.mp3')
            audio_files = sorted(glob.glob('memory/chunks/*.mp3'), key=os.path.getctime)
            transcripts = []
            for file in audio_files:
                transcripts.append(model.transcribe(file, language='en', fp16=False)['text'])
        
            messages = [
                {"role": "system", "content": "The transcripts are from a medical lecture. Create notes by extracting important information from the transcribed text. Correct any mispelled words. Format the notes in point form using - as the bullet point. Format notes using headings and subheadings as necessary. Put * around headings like *Pulmonary Stensosis* and ** around subheadings like **Diagnosis**. Subheadings may include summary, definitions, key terms, epidemiology, etiology, classification, pathophysiology, clinical features, diagnostics/investigations, differential diagnosis, treatment, prognosis, and complications."},
            ]
        elif fname.endswith('.pdf'):
            fname = fname.strip('.pdf')
            transcripts = pdf_to_content(f'static/upload/{fname}.pdf')
            messages = [
                {"role": "system", "content": "The content is from extracted PDF text from medical school. Create notes by extracting important information from the transcribed text. Correct any mispelled words. Format the notes in point form using - as the bullet point. Format notes using headings and subheadings as necessary. Put * around headings like *Pulmonary Stensosis* and ** around subheadings like **Diagnosis**. Subheadings may include summary, definitions, key terms, epidemiology, etiology, classification, pathophysiology, clinical features, diagnostics/investigations, differential diagnosis, treatment, prognosis, and complications."},
            ]
        elif fname.endswith('.mp3'):
            fname = fname.strip('.mp3')
            split_audio_audio(f'{fname}.mp3')
            audio_files = sorted(glob.glob('memory/chunks/*.mp3'), key=os.path.getctime)
            transcripts = []
            for file in tqdm.tqdm(audio_files):
                transcripts.append(model.transcribe(file, language='en', fp16=False)['text'])
        
            messages = [
                {"role": "system", "content": "The transcripts are from a medical lecture. Create notes by extracting important information from the transcribed text. Correct any mispelled words. Format the notes in point form using - as the bullet point. Format notes using headings and subheadings as necessary. Put * around headings like *Pulmonary Stensosis* and ** around subheadings like **Diagnosis**. Subheadings may include summary, definitions, key terms, epidemiology, etiology, classification, pathophysiology, clinical features, diagnostics/investigations, differential diagnosis, treatment, prognosis, and complications."},
            ]
        elif fname.endswith('.wav'):
            fname = fname.strip('.wav')
            split_audio_audio(f'{fname}.wav')
            audio_files = sorted(glob.glob('memory/chunks/*.mp3'), key=os.path.getctime)
            transcripts = []
            for file in audio_files:
                transcripts.append(model.transcribe(file, language='en', fp16=False)['text'])
        
            messages = [
                {"role": "system", "content": "The transcripts are from a medical lecture. Create notes by extracting important information from the transcribed text. Correct any mispelled words. Format the notes in point form using - as the bullet point. Format notes using headings and subheadings as necessary. Put * around headings like *Pulmonary Stensosis* and ** around subheadings like **Diagnosis**. Subheadings may include summary, definitions, key terms, epidemiology, etiology, classification, pathophysiology, clinical features, diagnostics/investigations, differential diagnosis, treatment, prognosis, and complications."},
            ]

        notes = []
        for transcript in tqdm.tqdm(transcripts):
            messages.append({"role": "user", "content": transcript})
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            notes.append(response['choices'][0]['message']['content'])
            messages.pop()

        with open('memory/notes.tkl', 'wb') as f:
            pickle.dump(notes, f)

    else:
        with open('memory/notes.tkl', 'rb') as f:
            notes = pickle.load(f)

    # saving notes
    with open(f'memory/{fname}.txt', 'w') as f:
        for note in notes:
            f.write(note)
            f.write('\n')

    doc = docx.Document()
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.15
    with open(f'memory/{fname}.txt') as f:
        for line in f:
            if line.startswith('*'):
                if line.startswith('**'):
                    # subheading
                    text = line.strip('**\n')
                    p = doc.add_heading(text, 3)
                else:
                    # heading
                    text = line.strip('*\n')
                    p = doc.add_heading(text, 2)
            else:
                # normal paragraph
                if line.startswith('\n'):
                    continue
                text = line.strip('-').strip(' -').strip()
                p = doc.add_paragraph(text, style='List Bullet')
    doc.save(f'static/upload/{fname}.docx')

    return f'Click to download: <a href="static/upload/{fname}.docx" target="_blank">{fname}.docx</a><br><br>If you would like to create questions from these notes:<br>Type any of:<br><br>1. Q&A: <em>!q number</em>. E.g., !q 5 will create 5 questions (max 10 questions per request).<br><br>2. MCQ: <em>!mcq number</em>. E.g., !mcq 3 will create 3 mcq questions (max 5 questions per request).'


def make_paper(fname, is_transcribed):
    if is_transcribed is False:
        fname = fname.strip('.pdf')
        transcripts = pdf_to_content(f'static/upload/{fname}.pdf')
        messages = [
            {"role": "system", "content": "The content is from extracted PDF text from a scientific paper. Create notes by extracting important information from the extracted text. Format the notes in point form using - as the bullet point. Format notes using headings and subheadings as necessary. Put * around headings like *Pulmonary Stensosis* and ** around subheadings like **Diagnosis**."},
        ]

        notes = []
        for transcript in transcripts:
            messages.append({"role": "user", "content": transcript})
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            notes.append(response['choices'][0]['message']['content'])
            messages.pop()

        with open('memory/notes.tkl', 'wb') as f:
            pickle.dump(notes, f)

    else:
        with open('memory/notes.tkl', 'rb') as f:
            notes = pickle.load(f)

    # saving notes
    with open(f'memory/{fname}.txt', 'w') as f:
        for note in notes:
            f.write(note)
            f.write('\n')

    doc = docx.Document()
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.15
    with open(f'memory/{fname}.txt') as f:
        for line in f:
            if line.startswith('*'):
                if line.startswith('**'):
                    # subheading
                    text = line.strip('**\n')
                    p = doc.add_heading(text, 3)
                else:
                    # heading
                    text = line.strip('*\n')
                    p = doc.add_heading(text, 2)
            else:
                # normal paragraph
                if line.startswith('\n'):
                    continue
                text = line.strip('-').strip(' -').strip()
                p = doc.add_paragraph(text, style='List Bullet')
    doc.save(f'static/upload/{fname}.docx')

    return f'Click to download: <a href="static/upload/{fname}.docx" target="_blank">{fname}.docx</a><br><br>If you would like to create questions from these notes:<br>Type any of:<br><br>1. Q&A: <em>!q number</em>. E.g., !q 5 will create 5 questions (max 10 questions per request).<br><br>2. MCQ: <em>!mcq number</em>. E.g., !mcq 3 will create 3 mcq questions (max 5 questions per request).'


def make_questions(npath, nquestions=5):
    try:
        nquestions = min(int(nquestions), 10)
    except:
        return "Please use integers for the number of questions you want."

    with open(npath, 'rb') as f:
        notes = pickle.load(f)

    random_note = random.choice(notes)
    messages = [
            {"role": "system", "content": f"This note is from a medical lecture. Create {nquestions} questions from this note with answers. Do not make up answers if you do not have the relevant context. Put each question and answer on a new line using <br>. Separate Q&A pairs by a new line using <br><br>. Format the questions using the following format: Q: <question><br>A: <answer><br><br>. For example, Q: What is the definition of pulmonary stenosis?<br>A: Narrowing of the pulmonary valve.<br><br>Q: ..."},
        ]

    messages.append({"role": "user", "content": random_note})
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )

    # saving conversation
    msg_path = 'memory/msg.tkl'
    if os.path.isfile(msg_path):
        with open(msg_path, 'rb') as f:
            messages = pickle.load(f)
    else:
        messages = [messages]

    messages.append({"role": "system", "content": response['choices'][0]['message']['content']})

    with open(msg_path, 'wb') as f:
        while len(messages) > 8:
            messages.pop(0)
        with open(msg_path, 'wb') as f:
            pickle.dump(messages, f)

    return response['choices'][0]['message']['content']


def make_mcqs(npath, nquestions=5):
    try:
        nquestions = min(int(nquestions), 5)
    except:
        return "Please use integers for the number of questions you want."

    with open(npath, 'rb') as f:
        notes = pickle.load(f)

    random_note = random.choice(notes)
    messages = [
            {"role": "system", "content": f"This note is from a medical lecture. Create {nquestions} single best answer multiple choice questions from this note with answers. Do not make up answers if you do not have the relevant context. Put each question and answer on a new line using <br>. Separate multiple choice pairs by a two new lines using <br><br>. For example: Q: What is the definition of pulmonary stenosis?<br>a) Narrowing of the pulmonary valve.<br>b) Degeneration of the pulmonary valve.<br>c) Absence of the pulmonary valve.<br>d) None of the above<br>Answer: a) Narrowing of the pulmonary valve.<br><br>Q: ..."},
        ]

    messages.append({"role": "user", "content": random_note})
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )

    # saving conversation
    msg_path = 'memory/msg.tkl'
    if os.path.isfile(msg_path):
        with open(msg_path, 'rb') as f:
            messages = pickle.load(f)
    else:
        messages = [messages]

    messages.append({"role": "system", "content": response['choices'][0]['message']['content']})

    with open(msg_path, 'wb') as f:
        while len(messages) > 8:
            messages.pop(0)
        with open(msg_path, 'wb') as f:
            pickle.dump(messages, f)

    return response['choices'][0]['message']['content']
